"""Build the client-facing QAOA RQP Pro V9.2 user manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "Version 9" / "docs"
ASSET_DIR = DOCS_DIR / "manual_assets"
OUTPUT_DOCX = DOCS_DIR / "QAOA_RQP_Pro_V9_2_Client_User_Manual.docx"
LOGO = ROOT / "assets" / "qubit-lab-light-blue-transparent.png"
WHITE_LOGO = ROOT / "assets" / "qubit-lab-logo-white.png"


# ─── Colour palette (mirrors qubit-lab.ch) ───────────────────────────────────
C = {
    "navy":       "0B1220",
    "navy_2":     "111827",
    "slate":      "1E293B",
    "slate_2":    "334155",
    "cyan":       "22D3EE",
    "cyan_dark":  "0891B2",
    "cyan_pale":  "CFFAFE",
    "amber":      "F59E0B",
    "amber_pale": "FEF3C7",
    "green":      "10B981",
    "green_pale": "D1FAE5",
    "red":        "EF4444",
    "red_pale":   "FEE2E2",
    "muted":      "94A3B8",
    "white":      "F8FAFC",
    "off_white":  "E2E8F0",
    "row_alt":    "F1F5F9",
    "ink":        "0F172A",
    "body":       "1E293B",
}


def rgb(hex_value: str) -> tuple[int, int, int]:
    h = hex_value.strip("#")
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def rgba(hex_value: str, a: int = 255) -> tuple[int, int, int, int]:
    return rgb(hex_value) + (a,)


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        ("/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold
         else "/System/Library/Fonts/Supplemental/Arial.ttf"),
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for p in candidates:
        try:
            return ImageFont.truetype(p, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy, width: int,
                  font_obj, fill, line_spacing: int = 8) -> int:
    x, y = xy
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textlength(candidate, font=font_obj) <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    for ln in lines:
        draw.text((x, y), ln, font=font_obj, fill=fill)
        y += font_obj.size + line_spacing
    return y


def _rrect(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def _paste_logo(img: Image.Image, path: Path, xy, max_width: int) -> None:
    if not path.exists():
        return
    logo = Image.open(path).convert("RGBA")
    scale = max_width / max(1, logo.width)
    logo = logo.resize((int(logo.width * scale), int(logo.height * scale)),
                       Image.LANCZOS)
    img.alpha_composite(logo, xy)


# ─── Image generators ────────────────────────────────────────────────────────

def _grid_background(draw: ImageDraw.ImageDraw, w: int, h: int) -> None:
    """Subtle diagonal grid in navy style."""
    for i in range(h):
        shade = int(max(0, min(255, 11 + i * 0.025)))
        draw.line((0, i, w, i), fill=(shade, 18, 34, 255))
    for x in range(-300, w + 100, 120):
        draw.line((x, h, x + 600, 0), fill=(34, 211, 238, 18), width=2)
    for y in range(100, h, 90):
        draw.line((0, y, w, y - 140), fill=(34, 211, 238, 12), width=1)


def create_cover_image() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "cover_rqp_manual.png"
    W, H = 1600, 900
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)

    # Cyan accent bar left
    draw.rectangle((0, 0, 6, H), fill=rgba(C["cyan"]))

    # Logo area
    _paste_logo(img, WHITE_LOGO if WHITE_LOGO.exists() else LOGO,
                (90, 68), 260)

    # Badge
    badge_text = "V 9.2"
    bw = int(draw.textlength(badge_text, font=_font(22, True))) + 30
    _rrect(draw, (90, 220, 90 + bw, 260), 8, rgba(C["cyan_dark"]))
    draw.text((105, 228), badge_text, font=_font(22, True),
              fill=rgb(C["white"]))

    # Title block
    draw.text((90, 276), "QAOA RQP Pro", font=_font(78, True),
              fill=rgb(C["white"]))
    draw.text((90, 370), "Rapid Quantum Prototyping for Portfolio Optimization",
              font=_font(36, True), fill=rgb(C["cyan"]))

    _draw_wrapped(
        draw,
        "End-to-end portfolio optimization using classical heuristics, "
        "QAOA simulation, and IBM quantum hardware — fully transparent, "
        "reproducible, and designed for business-facing review.",
        (92, 440), 960, _font(28), rgb(C["off_white"]), 10,
    )

    # Feature chips
    chips = [
        "Excel → QUBO", "Classical baseline", "QAOA Lightning / Tensor",
        "IBM Hardware", "Circuit exports",
    ]
    cx = 92
    for chip in chips:
        cw = int(draw.textlength(chip, font=_font(22, True))) + 44
        _rrect(draw, (cx, 700, cx + cw, 754), 22,
               rgba("082F49"), outline=rgb(C["cyan"]), width=2)
        draw.text((cx + 22, 716), chip, font=_font(22, True),
                  fill=rgb(C["cyan_pale"]))
        cx += cw + 16

    # Divider line
    draw.line((90, 790, 780, 790), fill=rgba(C["slate_2"]), width=2)
    draw.text((92, 806), "qubit-lab.ch  ·  Client User Manual  ·  Version 9.2.0",
              font=_font(22), fill=rgb(C["muted"]))

    # Right-side accent pattern
    for i, r in enumerate(range(120, 560, 80)):
        draw.ellipse((W - r - 20, H // 2 - r, W - 20, H // 2 + r),
                     outline=rgba(C["cyan_dark"], 30 - i * 5), width=2)

    img.convert("RGB").save(path, quality=96)
    return path


def create_workflow_image() -> Path:
    path = ASSET_DIR / "workflow_overview.png"
    W, H = 1600, 680
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)
    draw.rectangle((0, 0, 6, H), fill=rgba(C["cyan"]))

    draw.text((72, 50), "End-to-end workflow",
              font=_font(46, True), fill=rgb(C["white"]))
    draw.text((74, 108), "From Excel workbook to reproducible quantum-circuit exports",
              font=_font(26), fill=rgb(C["muted"]))

    steps = [
        ("01", C["cyan"],       "Prepare Workbook",
         "Assets, fixed/variable blocks, covariance, budgets, "
         "subtype constraints."),
        ("02", C["cyan_dark"],  "Inspect & Validate",
         "Structural checks, budget warnings, qubit count, "
         "and pre-run runtime estimate."),
        ("03", "F59E0B",        "Configure & Run",
         "Mode, worker profile, layers, iterations, restarts, "
         "second opinion."),
        ("04", C["green"],      "Monitor Live",
         "Backend logs, memory telemetry, phase progress, "
         "and live ETA."),
        ("05", "A78BFA",        "Review & Export",
         "Metrics, candidates, charts, QUBO breakdown, "
         "review JSON, notebooks."),
    ]

    step_w = 272
    gap = 30
    x0 = (W - (len(steps) * step_w + (len(steps) - 1) * gap)) // 2
    y0 = 178

    for i, (num, accent, title, desc) in enumerate(steps):
        x = x0 + i * (step_w + gap)
        # Card
        _rrect(draw, (x, y0, x + step_w, y0 + 370), 20, rgba(C["slate"]),
               outline=rgba(C["slate_2"]), width=2)
        # Number badge
        _rrect(draw, (x + 20, y0 + 20, x + 82, y0 + 82), 16,
               rgba(accent))
        draw.text((x + 51, y0 + 51), num, font=_font(28, True),
                  fill=rgb(C["navy"]), anchor="mm")
        # Title
        draw.text((x + 20, y0 + 102), title, font=_font(27, True),
                  fill=rgb(C["white"]))
        # Description
        _draw_wrapped(draw, desc, (x + 20, y0 + 154), step_w - 40,
                      _font(21), rgb(C["muted"]), 9)
        # Arrow
        if i < len(steps) - 1:
            ax = x + step_w + gap // 2
            ay = y0 + 185
            draw.line((ax - 10, ay, ax + 10, ay), fill=rgb(C["cyan_dark"]), width=4)
            draw.polygon([(ax + 10, ay), (ax, ay - 9), (ax, ay + 9)],
                         fill=rgb(C["cyan_dark"]))

    draw.text((72, H - 44), "qubit-lab.ch  ·  QAOA RQP Pro V9.2",
              font=_font(20), fill=rgb(C["muted"]))
    img.convert("RGB").save(path, quality=96)
    return path


def create_qaoa_circuit_image() -> Path:
    path = ASSET_DIR / "qaoa_circuit.png"
    W, H = 1600, 560
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)
    draw.rectangle((0, 0, 6, H), fill=rgba(C["amber"]))

    draw.text((72, 44), "QAOA circuit structure",
              font=_font(42, True), fill=rgb(C["white"]))
    draw.text((74, 98), "Standard x-mixer construction  ·  p layers alternating cost and mixer unitaries",
              font=_font(24), fill=rgb(C["muted"]))

    n_wires = 4
    wire_y = [200, 268, 336, 404]
    x_start = 120
    x_end = W - 80

    # Wire labels
    for i, y in enumerate(wire_y):
        draw.text((72, y - 14), f"q{i}", font=_font(22, True), fill=rgb(C["muted"]))
        draw.line((x_start, y, x_end, y), fill=rgba(C["slate_2"], 200), width=3)

    # |0⟩ labels
    for y in wire_y:
        draw.text((x_start - 56, y - 14), "|0⟩", font=_font(22), fill=rgb(C["cyan_pale"]))

    cur_x = x_start + 20

    # H gates
    for y in wire_y:
        _rrect(draw, (cur_x, y - 26, cur_x + 52, y + 26), 8, rgba(C["cyan_dark"]))
        draw.text((cur_x + 26, y), "H", font=_font(24, True),
                  fill=rgb(C["white"]), anchor="mm")
    draw.text((cur_x + 26, wire_y[-1] + 48), "init",
              font=_font(18), fill=rgb(C["muted"]), anchor="mm")
    cur_x += 72

    # Separator
    draw.line((cur_x + 10, 162, cur_x + 10, 428),
              fill=rgba(C["slate_2"], 140), width=2)
    cur_x += 26

    # Two cost+mixer layer blocks
    for layer in range(2):
        layer_x_start = cur_x

        # Cost layer: RZ terms
        for y in wire_y:
            _rrect(draw, (cur_x, y - 26, cur_x + 62, y + 26), 8, rgba(C["slate"]),
                   outline=rgba(C["amber"], 180), width=2)
            draw.text((cur_x + 31, y), "RZ", font=_font(22, True),
                      fill=rgb(C["amber"]), anchor="mm")
        cur_x += 80

        # CNOT pairs (ZZ interactions)
        for i in range(0, len(wire_y) - 1, 2):
            y1, y2 = wire_y[i], wire_y[i + 1]
            cx_x = cur_x + 26
            draw.line((cx_x, y1, cx_x, y2), fill=rgb(C["cyan"]), width=3)
            draw.ellipse((cx_x - 10, y1 - 10, cx_x + 10, y1 + 10),
                         fill=rgb(C["cyan"]))
            draw.ellipse((cx_x - 16, y2 - 16, cx_x + 16, y2 + 16),
                         outline=rgb(C["cyan"]), width=3)
            draw.line((cx_x - 14, y2, cx_x + 14, y2), fill=rgb(C["cyan"]), width=3)
            draw.line((cx_x, y2 - 14, cx_x, y2 + 14), fill=rgb(C["cyan"]), width=3)
        cur_x += 66

        # Mixer layer: RX terms
        for y in wire_y:
            _rrect(draw, (cur_x, y - 26, cur_x + 62, y + 26), 8, rgba(C["slate"]),
                   outline=rgba(C["green"], 180), width=2)
            draw.text((cur_x + 31, y), "RX", font=_font(22, True),
                      fill=rgb(C["green"]), anchor="mm")
        cur_x += 78

        # Layer brace
        layer_label = f"Layer {layer + 1}"
        brace_top = 152
        brace_bot = 432
        lx = layer_x_start - 6
        rx2 = cur_x - 8
        draw.line((lx, brace_top, rx2, brace_top), fill=rgba(C["muted"], 100), width=1)
        draw.line((lx, brace_bot, rx2, brace_bot), fill=rgba(C["muted"], 100), width=1)
        draw.text(((lx + rx2) // 2, brace_top - 26), layer_label,
                  font=_font(20, True), fill=rgb(C["muted"]), anchor="mm")

        # Separator between layers
        if layer == 0:
            draw.line((cur_x + 8, 162, cur_x + 8, 428),
                      fill=rgba(C["slate_2"], 120), width=1)
            cur_x += 22

    # Measure gates
    cur_x += 20
    for y in wire_y:
        _rrect(draw, (cur_x, y - 26, cur_x + 62, y + 26), 8, rgba("2D3748"),
               outline=rgba(C["muted"], 160), width=2)
        draw.text((cur_x + 31, y), "M", font=_font(22, True),
                  fill=rgb(C["muted"]), anchor="mm")

    # Legend
    lx = 900
    ly = H - 78
    items = [
        (C["amber"],    "Cost diagonal (RZ)"),
        (C["cyan"],     "ZZ entangling (CX-RZ-CX)"),
        (C["green"],    "Mixer (RX)"),
    ]
    for color, label in items:
        _rrect(draw, (lx, ly - 10, lx + 28, ly + 18), 4, rgba(color))
        draw.text((lx + 38, ly), label, font=_font(20), fill=rgb(C["muted"]))
        lx += int(draw.textlength(label, font=_font(20))) + 70

    draw.text((72, H - 44), "qubit-lab.ch  ·  QAOA RQP Pro V9.2",
              font=_font(20), fill=rgb(C["muted"]))
    img.convert("RGB").save(path, quality=96)
    return path


def create_access_levels_image() -> Path:
    path = ASSET_DIR / "access_levels.png"
    W, H = 1600, 680
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)
    draw.rectangle((0, 0, 6, H), fill=rgba(C["green"]))

    draw.text((72, 44), "Access levels and limits",
              font=_font(42, True), fill=rgb(C["white"]))
    draw.text((74, 98), "License key controls mode availability, qubit capacity, worker profiles, and export features",
              font=_font(24), fill=rgb(C["muted"]))

    levels = [
        ("Public Demo",     "0", C["muted"],      "8",  "1 / 10 / 1",   "60 s",   "5 MB",   "—",      "—"),
        ("Qualified Demo",  "1", C["cyan_dark"],   "16", "2 / 20 / 1",   "4 min",  "10 MB",  "—",      "—"),
        ("Tester",          "2", C["cyan"],        "24", "6 / 200 / 3",  "45 min", "25 MB",  "✓",      "✓"),
        ("Internal Power",  "3", C["amber"],       "24", "8 / 300 / 5",  "2 h",    "25 MB",  "✓",      "✓"),
        ("Internal Ultra",  "5", C["green"],       "35", "10 / 300 / 3", "30 d",   "25 MB",  "✓",      "✓"),
    ]

    col_labels = ["Level", "ID", "QAOA qubits", "Layers/Iter/Restarts",
                  "Max runtime", "Upload", "IBM HW", "Exports"]
    col_x = [72, 280, 400, 590, 870, 1040, 1200, 1360]
    header_y = 166

    # Header row
    _rrect(draw, (60, header_y - 8, W - 60, header_y + 40), 8, rgba(C["slate"]))
    for label, x in zip(col_labels, col_x):
        draw.text((x, header_y + 8), label, font=_font(19, True), fill=rgb(C["cyan"]))

    row_y = header_y + 54
    for i, (name, lid, accent, qubits, layers_etc, runtime, upload, ibm, exports) in enumerate(levels):
        row_bg = C["slate"] if i % 2 == 0 else "1A2640"
        _rrect(draw, (60, row_y - 6, W - 60, row_y + 42), 6, rgba(row_bg))
        # Accent dot
        draw.ellipse((66, row_y + 12, 82, row_y + 28), fill=rgba(accent))
        draw.text((col_x[0] + 4, row_y + 8), name, font=_font(21, True), fill=rgb(C["white"]))
        draw.text((col_x[1], row_y + 8), lid, font=_font(21), fill=rgb(C["muted"]))
        draw.text((col_x[2], row_y + 8), qubits, font=_font(21), fill=rgb(C["white"]))
        draw.text((col_x[3], row_y + 8), layers_etc, font=_font(19), fill=rgb(C["off_white"]))
        draw.text((col_x[4], row_y + 8), runtime, font=_font(19), fill=rgb(C["off_white"]))
        draw.text((col_x[5], row_y + 8), upload, font=_font(19), fill=rgb(C["muted"]))
        draw.text((col_x[6], row_y + 8), ibm,
                  font=_font(21, True),
                  fill=rgb(C["green"] if ibm == "✓" else C["muted"]))
        draw.text((col_x[7], row_y + 8), exports,
                  font=_font(21, True),
                  fill=rgb(C["green"] if exports == "✓" else C["muted"]))
        row_y += 54

    draw.text((72, H - 44), "qubit-lab.ch  ·  QAOA RQP Pro V9.2",
              font=_font(20), fill=rgb(C["muted"]))
    img.convert("RGB").save(path, quality=96)
    return path


def create_results_overview_image() -> Path:
    path = ASSET_DIR / "results_overview.png"
    W, H = 1600, 820
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)
    draw.rectangle((0, 0, 6, H), fill=rgba(C["green"]))

    draw.text((72, 44), "Results dashboard overview",
              font=_font(42, True), fill=rgb(C["white"]))
    draw.text((74, 98), "Key output panels and how to interpret them",
              font=_font(24), fill=rgb(C["muted"]))

    panels = [
        # (x, y, w, h, accent, title, lines)
        (72,  158, 460, 280, C["cyan"],    "Metric Cards (×3)",
         ["Classical / Quantum / IBM Hardware", "QUBO value, selected amount,",
          "return, volatility, Sharpe ratio", "Budget gap and probability"]),
        (562, 158, 460, 280, C["amber"],   "Candidate Tables",
         ["Top classical candidates (QUBO rank)", "Top quantum by QUBO value",
          "Top quantum by probability", "IBM hardware QUBO + probability"]),
        (1052, 158, 476, 280, C["green"],  "QUBO Breakdown",
         ["Return, risk, budget, type-budget", "Per-term contribution",
          "Compare classical vs. quantum", "Diagnose penalty balance"]),
        (72,  476, 460, 280, "A78BFA",    "Optimization History",
         ["Expected QUBO energy vs. iteration", "Convergence across restarts",
          "Best energy trajectory", "Detect local-minima trapping"]),
        (562, 476, 460, 280, C["cyan_dark"], "Type-Budget Diagnostics",
         ["Achieved vs. target per subtype", "Relative deviation %",
          "Penalty contribution per term", "V9.2 addition — up to 5 types"]),
        (1052, 476, 476, 280, C["muted"],  "Backend Logs & Diagnostics",
         ["Live timestamped execution log", "Export cap diagnostics",
          "IBM hardware job details", "Memory telemetry over time"]),
    ]

    for px, py, pw, ph, accent, title, lines in panels:
        _rrect(draw, (px, py, px + pw, py + ph), 16, rgba(C["slate"]),
               outline=rgba(C["slate_2"]), width=2)
        # Accent top bar
        draw.rectangle((px, py, px + pw, py + 6), fill=rgba(accent))
        draw.text((px + 20, py + 22), title, font=_font(24, True), fill=rgb(accent))
        for i, line in enumerate(lines):
            draw.text((px + 20, py + 68 + i * 38), "·  " + line,
                      font=_font(20), fill=rgb(C["muted"]))

    draw.text((72, H - 44), "qubit-lab.ch  ·  QAOA RQP Pro V9.2",
              font=_font(20), fill=rgb(C["muted"]))
    img.convert("RGB").save(path, quality=96)
    return path


def create_exports_image() -> Path:
    path = ASSET_DIR / "exports_panel.png"
    W, H = 1600, 600
    img = Image.new("RGBA", (W, H), rgba(C["navy"]))
    draw = ImageDraw.Draw(img)
    _grid_background(draw, W, H)
    draw.rectangle((0, 0, 6, H), fill=rgba(C["amber"]))

    draw.text((72, 44), "Review files and code exports",
              font=_font(42, True), fill=rgb(C["white"]))
    _draw_wrapped(
        draw,
        "All optimization artifacts can be saved as a review JSON and reloaded "
        "without rerunning the backend. The same result package generates "
        "executable circuits in Qiskit, PennyLane, and Cirq.",
        (74, 100), 1100, _font(24), rgb(C["muted"]), 10,
    )

    # Review file panel
    _rrect(draw, (72, 198, 500, 548), 20, rgba(C["slate"]),
           outline=rgba(C["slate_2"]), width=2)
    draw.text((100, 224), "Review File", font=_font(28, True), fill=rgb(C["cyan"]))
    for i, btn in enumerate(["Save Review File", "Download Raw JSON Data",
                              "Load Review / Raw JSON File"]):
        by = 280 + i * 74
        _rrect(draw, (100, by, 472, by + 52), 12, rgba(C["cyan_dark"]),
               outline=rgba(C["cyan"]), width=2)
        draw.text((120, by + 14), btn, font=_font(22, True), fill=rgb(C["white"]))

    # Code exports panel
    _rrect(draw, (556, 198, W - 72, 548), 20, rgba(C["slate"]),
           outline=rgba(C["slate_2"]), width=2)
    draw.text((584, 224), "Code Exports", font=_font(28, True), fill=rgb(C["amber"]))
    exports = [
        ("Qiskit Notebook",    ".ipynb", "QuantumCircuit + statevector, IBM Runtime cell"),
        ("Qiskit Python",       ".py",   "Script-style reproducible review"),
        ("PennyLane Notebook", ".ipynb", "default.qubit simulation, Python 3.11+"),
        ("Google Cirq Notebook",".ipynb","Cirq simulator comparison"),
    ]
    for i, (name, fmt, desc) in enumerate(exports):
        ey = 280 + i * 64
        fw = int(draw.textlength(fmt, font=_font(19, True))) + 22
        _rrect(draw, (584, ey, 584 + fw, ey + 34), 8, rgba(C["slate_2"]))
        draw.text((594, ey + 7), fmt, font=_font(19, True), fill=rgb(C["amber"]))
        draw.text((584 + fw + 16, ey + 6), name, font=_font(21, True),
                  fill=rgb(C["white"]))
        draw.text((584 + fw + 16, ey + 32), desc, font=_font(18),
                  fill=rgb(C["muted"]))

    draw.text((72, H - 44), "qubit-lab.ch  ·  QAOA RQP Pro V9.2",
              font=_font(20), fill=rgb(C["muted"]))
    img.convert("RGB").save(path, quality=96)
    return path


# ─── docx helpers ────────────────────────────────────────────────────────────

def _shd(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), fill.upper())
    s.set(qn("w:color"), "auto")
    s.set(qn("w:val"), "clear")
    tc_pr.append(s)


def _cell_color(cell, hex_color: str, bold: bool = False) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor.from_string(hex_color)
            if bold:
                run.font.bold = True


def _hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), C["cyan_dark"])
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def _set_styles(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Aptos"
    n.font.size = Pt(10.5)
    n.font.color.rgb = RGBColor.from_string(C["body"])

    for style_name, size, color in [
        ("Title",     28, C["navy"]),
        ("Heading 1", 18, C["navy"]),
        ("Heading 2", 13, C["cyan_dark"]),
        ("Heading 3", 11, C["slate"]),
    ]:
        s = doc.styles[style_name]
        s.font.name = "Aptos Display"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True


def _section_banner(doc: Document, text: str,
                    bg: str = C["navy"], fg: str = C["cyan"],
                    accent_bar: str = C["cyan"]) -> None:
    """Full-width dark banner for major section headers."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _shd(cell, bg)
    cell.width = Inches(7.3)
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "160" if side in ("top", "bottom") else "200")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(fg)
    doc.add_paragraph()


def _callout(doc: Document, title: str, body: str,
             bg: str = "E0F2FE", title_color: str = C["navy"]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _shd(cell, bg)
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "140")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(title_color)
    p.add_run(body)
    doc.add_paragraph()


def _dark_callout(doc: Document, label: str, body: str,
                  label_color: str = C["cyan"]) -> None:
    """Dark navy callout for tips/notes."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _shd(cell, C["navy_2"])
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "140")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)
    bdr = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), label_color)
    bdr.append(left)
    tc_pr.append(bdr)
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(label_color)
    t = p.add_run(body)
    t.font.color.rgb = RGBColor.from_string(C["off_white"])
    doc.add_paragraph()


def _table(doc: Document, headers: list[str], rows: list[list[str]],
           widths: list[float] | None = None,
           alt_rows: bool = True) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _shd(hdr_cells[i], C["navy"])
        _cell_color(hdr_cells[i], "FFFFFF", bold=True)
        if widths:
            hdr_cells[i].width = Cm(widths[i])
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            if alt_rows and ri % 2 == 1:
                _shd(cells[i], C["row_alt"])
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _img(doc: Document, path: Path, width: float = 7.3) -> None:
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _spacer(doc: Document) -> None:
    doc.add_paragraph()


# ─── Document build ───────────────────────────────────────────────────────────

def build_doc() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    print("Generating images…")
    cover    = create_cover_image()
    workflow = create_workflow_image()
    circuit  = create_qaoa_circuit_image()
    access   = create_access_levels_image()
    results  = create_results_overview_image()
    exports  = create_exports_image()
    print("Images done.")

    doc = Document()
    _set_styles(doc)

    sec = doc.sections[0]
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(1.9)
    sec.right_margin  = Cm(1.9)

    hdr = sec.header.paragraphs[0]
    hdr.text = "QAOA RQP Pro V9.2  ·  Client User Manual  ·  qubit-lab.ch"
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hdr.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(C["muted"])

    ftr = sec.footer.paragraphs[0]
    ftr.text = "© qubit-lab.ch  ·  Version 9.2.0  ·  For authorized use only"
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ftr.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(C["muted"])

    # ── Cover ──────────────────────────────────────────────────────────────
    _img(doc, cover, 7.3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Client User Manual")
    r.bold = True
    r.font.size = Pt(14)
    p2 = doc.add_paragraph(
        "Inputs, methods, outputs, interpretation, and reproducible exports"
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacer(doc)
    _callout(
        doc,
        "Positioning.",
        "QAOA RQP Pro is a rapid prototyping environment for portfolio-optimization "
        "workflows. It converts Excel workbook inputs into classical baselines, QUBO "
        "formulations, QAOA simulations, IBM hardware runs, live diagnostics, and "
        "portable quantum-circuit exports — all in a single browser-based cockpit.",
        "E0F2FE", C["navy"],
    )

    doc.add_page_break()

    # ── Contents ───────────────────────────────────────────────────────────
    doc.add_heading("Contents", level=1)
    _bullets(doc, [
        "1.  About qubit-lab.ch and the RQP mission",
        "2.  What the tool does",
        "3.  Quick start",
        "4.  Inputs and workbook format",
        "5.  Optimization settings",
        "6.  Methods used",
        "7.  Execution and live progress",
        "8.  Outputs and interpretation",
        "9.  IBM Quantum Hardware second opinion",
        "10. Review files and code exports",
        "11. Access levels and limits",
        "12. Worker profiles",
        "13. Memory telemetry",
        "14. Good practice and troubleshooting",
        "Appendix A — Settings field reference",
        "Appendix B — Glossary",
    ])

    doc.add_page_break()

    # ── 1. About ───────────────────────────────────────────────────────────
    _section_banner(doc, "1.  About qubit-lab.ch and the RQP Mission")
    p = doc.add_paragraph()
    p.add_run("qubit-lab.ch").bold = True
    p.add_run(
        " bridges quantum computing research and real-world applications in business "
        "and finance. Its practical mission is to make quantum methods transparent, "
        "understandable, and testable for finance professionals, technology stakeholders, "
        "quants, and technically curious practitioners — without hiding the mechanics "
        "inside a black box."
    )
    doc.add_paragraph(
        "The Rapid Quantum Prototyping (RQP) tool embodies this mission. It takes a "
        "familiar Excel-based portfolio use case and exposes every step of the pipeline: "
        "workbook parsing, QUBO construction, classical baseline, QAOA simulation, "
        "live diagnostics, result interpretation, and portable circuit exports. Users "
        "can inspect, challenge, and reproduce every component."
    )
    _dark_callout(
        doc,
        "Straight talk, not black-box quantum.",
        "The full optimization chain is always visible: workbook assumptions, QUBO "
        "construction math, solver settings, backend logs, portfolio metrics, candidate "
        "tables, charts, IBM hardware diagnostics, and executable circuit notebooks.",
        C["cyan"],
    )
    p3 = doc.add_paragraph()
    p3.add_run("Further reading:  ")
    _hyperlink(p3, "qubit-lab.ch", "https://qubit-lab.ch/")
    p3.add_run("  ·  ")
    _hyperlink(p3, "Quantum Finance / RQP", "https://qubit-lab.ch/finance")

    # ── 2. What the tool does ──────────────────────────────────────────────
    _section_banner(doc, "2.  What the Tool Does")
    doc.add_paragraph(
        "QAOA RQP Pro V9.2 is a browser-based optimization cockpit for portfolio "
        "experiments. Users upload a structured Excel workbook, configure optimization "
        "settings, and submit a backend job that evaluates a classical heuristic "
        "baseline and — when selected — a QAOA simulation path. After the job "
        "finishes, the interface provides portfolio metrics, candidate tables, charts, "
        "diagnostics, and downloadable artifacts."
    )
    _bullets(doc, [
        "Parses a portfolio workbook (fixed holdings, variable candidates, covariance, "
        "returns, costs, settings) into a binary optimization problem.",
        "Separates fixed positions (always included, no qubit consumed) from variable "
        "decision blocks (each consumes one qubit).",
        "Constructs a QUBO/Ising representation incorporating return, variance/risk, "
        "total budget, and up to five exact subtype budget constraints — new in V9.",
        "Runs a classical heuristic candidate search and, when selected, a QAOA "
        "simulation on PennyLane's lightning.qubit or default.tensor backend.",
        "Provides an optional second-opinion comparison via Qiskit simulation or "
        "real IBM quantum hardware execution.",
        "Streams live backend logs, ETA, memory telemetry, and phase progress "
        "during job execution.",
        "Exports completed results as review JSON and executable Qiskit, PennyLane, "
        "and Cirq notebook files.",
    ])
    _img(doc, workflow)

    doc.add_page_break()

    # ── 3. Quick start ─────────────────────────────────────────────────────
    _section_banner(doc, "3.  Quick Start")
    _numbered(doc, [
        "Open the RQP tool from qubit-lab.ch → Finance / RQP, or from a private "
        "V9.2 route if provided.",
        "Paste a license key into the Access panel and click Check License, or "
        "click Check License or Public Demo to proceed under public demo limits.",
        "Download a demo workbook (7, 16, or 24 qubits) using the buttons at the "
        "top, or upload your own .xlsx file using the file selector.",
        "Review the Workbook Summary — it confirms qubit count, QUBO shape, "
        "fixed/variable split, type constraints, and loaded settings.",
        "In Optimization Settings, choose mode, worker profile, response level, "
        "QAOA parameters, and optional second opinion.",
        "If using IBM Hardware, expand the IBM Quantum session panel, paste your "
        "IBM Quantum token, and review the pre-run hardware depth estimate.",
        "Click Run Optimization on V9.2 and monitor the Backend Optimization Log, "
        "memory telemetry, and live ETA.",
        "When the job finishes, review metric cards, portfolio contents, candidate "
        "tables, charts, QUBO breakdown, and diagnostics.",
        "Save a Review File or Raw JSON for audit and reproducibility. Use Code "
        "Exports to generate Qiskit, PennyLane, or Cirq notebooks.",
    ])
    _dark_callout(
        doc,
        "Tip.",
        "Always inspect the workbook before running. The inspect step is fast, validates "
        "structure, flags budget issues, counts qubits, and shows a runtime estimate — "
        "helping you catch misconfiguration before committing a run.",
        C["amber"],
    )

    # ── 4. Inputs ──────────────────────────────────────────────────────────
    _section_banner(doc, "4.  Inputs and Workbook Format")
    doc.add_paragraph(
        "The primary input is an Excel workbook (.xlsx). Version 9 extends the "
        "Version 8 format with optional exact subtype budget constraints. Three sheets "
        "are required; additional sheets are recognized but optional."
    )
    doc.add_heading("4.1  Required Sheets", level=2)
    _table(
        doc,
        ["Sheet", "Purpose", "Key fields"],
        [
            ["Settings", "Global optimization and portfolio configuration.",
             "Budget, Lambda Budget, Lambda Variance, Risk Free Rate, QAOA P, "
             "QAOA Maxiter, QAOA Multistart Restarts, QAOA Shots, Warm Start, "
             "Restart Perturbation, Random Seed, Additional Type Constraints, "
             "Type A–E Name / Budget / Budget Penalty."],
            ["Assets", "One row per portfolio block (fixed or variable).",
             "Ticker, Company, Decision Role, Option, Indicative Market Cost USD, "
             "Expected Return Proxy, Annual Volatility, and optional "
             "Type A Size … Type E Size columns."],
            ["AnnualizedCovariance", "Annualized covariance matrix for all tickers.",
             "Square matrix; rows and columns must match ticker labels in Assets."],
        ],
        [3.2, 4.8, 8.4],
    )
    doc.add_heading("4.2  Asset Rows", level=2)
    _table(
        doc,
        ["Decision Role", "Meaning", "Qubit consumed?"],
        [
            ["fixed", "Always included in the portfolio. Contributes cost, return, "
             "risk, budget usage, and type-budget exposure as a constant offset.", "No"],
            ["variable", "Binary decision variable. The optimizer chooses whether "
             "to include this block.", "Yes — one qubit per row."],
        ],
        [2.8, 10.0, 3.6],
    )
    _bullets(doc, [
        "Indicative Market Cost USD is the primary cost column in V9. A legacy "
        "Approx Cost USD column is mapped for compatibility; when both are present, "
        "the indicative column takes priority.",
        "Expected Return Proxy and Annual Volatility must be numeric for variable "
        "rows. Missing or non-numeric values cause a validation error.",
        "Fixed rows must have numeric costs and returns (they influence the constant "
        "offset of the QUBO and the Sharpe-ratio metrics).",
    ])
    doc.add_heading("4.3  Workbook Validation Warnings", level=2)
    doc.add_paragraph(
        "The inspect step runs automatic sanity checks and surfaces warnings if any "
        "of the following conditions are detected:"
    )
    _bullets(doc, [
        "Fixed holdings exceed the configured total budget before any variable block "
        "is selected.",
        "Budget is smaller than fixed holdings plus the cheapest variable block.",
        "Fixed holdings plus the median variable block cost already exceed the budget.",
        "No variable decision blocks found — optimizer has nothing to select.",
        "Available variable candidate universe is too small to reach the configured "
        "budget.",
    ])
    _callout(
        doc,
        "Note.",
        "Warnings do not block a run, but they should be reviewed. A run with "
        "unresolved budget warnings will likely produce low-quality or trivially "
        "identical candidates.",
        "FEF3C7", C["navy"],
    )
    doc.add_heading("4.4  Optional Subtype Budget Constraints (V9 Addition)", level=2)
    doc.add_paragraph(
        "V9.2 supports up to five exact subtype budget targets layered on top of "
        "the main total-budget constraint. These can represent any meaningful "
        "portfolio dimension: sector, region, asset class, sustainability bucket, "
        "liquidity tier, or internal risk category."
    )
    _table(
        doc,
        ["Settings-sheet field", "Description"],
        [
            ["Additional Type Constraints", "Integer 0–5. Missing or blank is treated as 0."],
            ["Type X Name", "User-facing label (e.g., Compute Infra). Internal IDs "
             "type_a … type_e are stable and independent of this label."],
            ["Type X Budget", "Target exposure for that subtype (positive number, "
             "same currency as main budget)."],
            ["Type X Budget Penalty", "Penalty weight λ_k applied to deviations from "
             "the subtype target. Must be ≥ 0."],
            ["Type A Size … Type E Size (Assets sheet)", "Numeric exposure of each "
             "asset block to the subtype. Missing cells are treated as 0. Non-empty "
             "non-numeric values fail validation. Do not rename these columns."],
        ],
        [5.4, 11.0],
    )
    _dark_callout(
        doc,
        "Constraint stability.",
        "The internal constraint IDs type_a … type_e are fixed regardless of the "
        "user-facing name. Rename the Type X Name label freely; never rename the "
        "Type A Size column headers — the parser matches them exactly.",
        C["cyan"],
    )

    doc.add_page_break()

    # ── 5. Optimization settings ───────────────────────────────────────────
    _section_banner(doc, "5.  Optimization Settings")
    doc.add_paragraph(
        "All settings can be provided from three sources in priority order: "
        "(1) the UI form, (2) the workbook Settings sheet, (3) backend defaults. "
        "The diagnostics block in the completed result shows which source was "
        "effective for each setting."
    )
    doc.add_heading("5.1  Mode", level=2)
    _table(
        doc,
        ["Mode", "Internal value", "Description"],
        [
            ["Classical Only", "classical_only",
             "Runs the heuristic classical candidate search only. No QAOA circuit "
             "is built or executed. Fast; use for workbook validation."],
            ["QAOA Lightning Sim", "qaoa_lightning_sim",
             "PennyLane lightning.qubit — high-performance C++ statevector simulator. "
             "For ≤ 24 qubits, runs in exact-probability mode (full statevector, "
             "no sampling noise). Above 24 qubits, switches to sampling."],
            ["QAOA Tensor Sim", "qaoa_tensor_sim",
             "PennyLane default.tensor — tensor-network simulator. Always sampling. "
             "Suitable for larger qubit counts where full statevector is impractical."],
        ],
        [3.4, 3.6, 9.4],
    )
    doc.add_heading("5.2  Second Opinion / Comparison", level=2)
    _table(
        doc,
        ["Option", "Internal value", "Description", "Level required"],
        [
            ["Internal only", "internal_only",
             "No second opinion. Single QAOA result.", "All"],
            ["Qiskit simulation", "qiskit_export",
             "Reconstructs the optimized circuit in Qiskit and evaluates it "
             "as a statevector simulation. Cross-framework validation.", "Tester+"],
            ["Qiskit on IBM Hardware", "ibm_external_run",
             "Transpiles and submits to a real IBM quantum device. "
             "Requires an IBM Quantum token.", "Tester+"],
        ],
        [3.2, 3.2, 6.6, 2.8],
    )
    doc.add_heading("5.3  Response Level", level=2)
    _table(
        doc,
        ["Level", "Content", "Use case"],
        [
            ["compact", "Summary cards only.",
             "Quick status check; minimal data transfer."],
            ["standard", "Summary cards plus key tables.",
             "General result review."],
            ["full", "All blocks including backend diagnostics, export diagnostics, "
             "memory telemetry, IBM hardware diagnostics, circuit overview, "
             "solver comparison.",
             "Client review, documentation, and technical audit. Recommended."],
        ],
        [2.4, 6.4, 7.6],
    )
    doc.add_heading("5.4  QAOA Parameters", level=2)
    _table(
        doc,
        ["Parameter", "Workbook key", "Description", "Guidance"],
        [
            ["Layers (p)", "QAOA P",
             "QAOA depth — number of alternating cost + mixer layers.",
             "Start at 1–2 for demos. Higher depth can improve quality but "
             "increases runtime and tuning complexity."],
            ["Iterations", "QAOA Maxiter",
             "Maximum optimizer iterations per restart.",
             "50–100 for demos; 150–300 for deeper searches."],
            ["Restarts", "QAOA Multistart Restarts",
             "Independent optimizer starts from fresh (perturbed) initial angles.",
             "1 is fast; 3+ helps avoid local minima, especially at higher depth."],
            ["QAOA Shots", "QAOA Shots",
             "Readout samples when sampling mode is active.",
             "Shown as exact when exact-probability mode is used (Lightning Sim ≤ 24 qubits)."],
            ["Warm Start", "Warm Start",
             "Initializes layer p angles from optimized layer p−1 solution.",
             "Adds ~50% runtime but often produces better starting points."],
            ["Restart Perturbation", "Restart Perturbation",
             "Magnitude of random noise added to initial angles at each restart.",
             "Default 0.05. Increase if restarts converge too similarly."],
            ["Budget Lambda", "Lambda Budget",
             "Penalty weight enforcing the total budget constraint.",
             "Start at 50. Increase (100–200) for stricter budget adherence."],
            ["Risk Lambda", "Lambda Variance",
             "Penalty weight on portfolio variance/risk.",
             "Start at 6. Increase for lower-risk portfolios."],
            ["Risk-Free Rate", "Risk Free Rate",
             "Annual reference rate for Sharpe-like metric calculation.",
             "E.g., 0.04 for 4%. Match to return horizon."],
            ["Random Seed", "Random Seed",
             "Optional integer 0–4,294,967,295 for reproducible initialization.",
             "Same seed + same code + environment → same result. "
             "Not guaranteed across version changes."],
        ],
        [2.8, 3.4, 4.8, 5.4],
    )

    doc.add_page_break()

    # ── 6. Methods ─────────────────────────────────────────────────────────
    _section_banner(doc, "6.  Methods Used")
    doc.add_heading("6.1  Problem Formulation", level=2)
    doc.add_paragraph(
        "The optimizer treats each variable asset block as a binary decision variable "
        "x_i ∈ {0, 1}. Fixed positions are always included and contribute constant "
        "offsets to the objective. The goal is to find the assignment of x_i values "
        "that minimizes the combined objective."
    )
    doc.add_heading("6.2  QUBO Construction", level=2)
    doc.add_paragraph(
        "The full QUBO (Quadratic Unconstrained Binary Optimization) objective combines "
        "four components:"
    )
    _table(
        doc,
        ["Term", "Formula", "Role"],
        [
            ["Return",
             "−λ_r · Σ_i w_i · x_i",
             "Rewards higher expected-return portfolios (negative contribution lowers QUBO)."],
            ["Variance / Risk",
             "λ_v · x^T Σ x",
             "Penalizes portfolio variance scaled by the risk lambda."],
            ["Budget",
             "λ_b · ((Σ_i c_i · x_i + C_fixed) / B − 1)²",
             "Penalizes deviation from the total budget target B."],
            ["Subtype budgets (×k)",
             "λ_k · ((E_k^fixed + Σ_i s_ki · x_i) / B_k − 1)²",
             "One term per active type constraint. Penalizes deviation from "
             "each subtype exposure target B_k."],
        ],
        [2.6, 6.6, 7.2],
    )
    doc.add_paragraph(
        "In the subtype term, E_k^fixed is the sum of Type X Size values across all "
        "fixed blocks (a constant offset), s_ki is the Type X Size for variable block i, "
        "and the normalized form (achieved/budget − 1)² means a penalty of zero when "
        "the target is hit exactly."
    )
    _dark_callout(
        doc,
        "Why QUBO matters.",
        "QUBO is the bridge between the business problem and the quantum circuit. It "
        "maps every candidate portfolio to a single bitstring and assigns it an "
        "objective value that can be searched, ranked, compared, and reported.",
        C["amber"],
    )
    doc.add_heading("6.3  Ising Conversion", level=2)
    doc.add_paragraph(
        "The QUBO matrix Q is converted to Ising form for quantum circuit construction "
        "using the standard substitution x_i = (1 − σ_z^i) / 2. This produces diagonal "
        "Ising h coefficients (single-spin terms) and off-diagonal J coefficients "
        "(spin-spin interaction terms). The Ising Hamiltonian H_C is the QAOA cost "
        "operator applied in the cost unitary U_C(γ)."
    )
    doc.add_heading("6.4  Classical Baseline", level=2)
    doc.add_paragraph(
        "The classical path uses a heuristic search combining random sampling and local "
        "improvement over the binary decision space. It evaluates candidate portfolios "
        "under the same QUBO objective and exports the top unique candidates ranked by "
        "QUBO value. The classical baseline is always run regardless of mode and "
        "provides a practical benchmark for quantum results."
    )
    doc.add_paragraph(
        "The candidate count exported is capped at the number of unique candidates "
        "found after deduplication and search convergence. The diagnostic field "
        "classical_export_cap_reason explains any reduction."
    )
    doc.add_heading("6.5  QAOA Circuit and Simulation", level=2)
    doc.add_paragraph(
        "The QAOA circuit uses the standard transverse-field mixer construction:"
    )
    _bullets(doc, [
        "Initialization: Hadamard gates applied to all n qubits, creating a uniform "
        "superposition over all 2^n possible bitstrings.",
        "Cost unitary U_C(γ): RZ(2γ·h_i) rotations for single-qubit diagonal terms; "
        "CX → RZ(2γ·J_ij) → CX sequences for each two-qubit ZZ interaction term.",
        "Mixer unitary U_B(β): RX(2β) applied to every qubit (transverse-field mixer).",
        "Layers: cost + mixer repeated p times with independently optimized angles "
        "γ = (γ_1, …, γ_p) and β = (β_1, …, β_p).",
    ])
    _img(doc, circuit)
    _table(
        doc,
        ["Gate", "Count (n qubits, p layers, k ZZ terms)"],
        [
            ["H (Hadamard, init)", "n"],
            ["RZ (cost diagonal)", "(n + k) · p"],
            ["CX (CNOT, ZZ coupling)", "2 · k · p"],
            ["RX (mixer)", "n · p"],
            ["Total 1Q gates", "n + (n + k + n) · p = n(1 + 2p) + k·p"],
            ["Total 2Q gates", "2 · k · p"],
        ],
        [6.2, 10.2],
    )
    doc.add_paragraph(
        "For the 16-qubit demo run (p=2, k=120): 800 total gates, 480 CX gates, "
        "320 single-qubit gates. After transpilation to IBM ibm_fez native gates: "
        "4,831 total transpiled gates, 1,158 CZ gates, transpiled depth 963."
    )
    doc.add_heading("6.6  Exact Probability vs. Sampling Mode", level=2)
    _table(
        doc,
        ["Condition", "Mode", "Behavior"],
        [
            ["Lightning Sim, ≤ 24 qubits", "Exact probability",
             "Full statevector computed. Every 2^n bitstring receives an exact "
             "probability. No shot noise. Typical for ≤ 16-qubit demos."],
            ["Lightning Sim, > 24 qubits", "Sampling",
             "qaoa_shots measurements drawn from the statevector probability "
             "distribution. Probabilities estimated from observed frequencies."],
            ["Tensor Sim (any qubit count)", "Sampling",
             "Tensor-network backend always uses sampling — full statevector is "
             "not computed explicitly."],
        ],
        [3.8, 3.4, 9.2],
    )
    doc.add_heading("6.7  Parameter Optimization and Multi-start", level=2)
    doc.add_paragraph(
        "The QAOA angle parameters γ and β are optimized using a classical "
        "gradient-free optimizer that minimizes the expected QUBO energy "
        "⟨ψ(γ,β)|H_C|ψ(γ,β)⟩. Each restart begins from a fresh initialization "
        "perturbed by the restart_perturbation magnitude. The optimizer returns "
        "best_gammas and best_betas — the angle arrays that achieved the lowest "
        "energy across all restarts and iterations. These are embedded in the result "
        "JSON and used for all code exports."
    )
    _callout(
        doc,
        "Warm start.",
        "When warm_start = true, the initial angles for layer p+1 are seeded from "
        "the optimized layer p solution rather than random initialization. This "
        "typically improves convergence at the cost of ~50% additional runtime.",
        "E0F2FE", C["navy"],
    )
    doc.add_heading("6.8  Runtime Estimation", level=2)
    doc.add_paragraph(
        "The pre-run estimate uses the following calibrated formulas:"
    )
    _table(
        doc,
        ["Estimator", "Formula"],
        [
            ["QAOA raw estimate",
             "safety_factor × α × state_count × layers × iterations × restarts "
             "× (1.5 if warm_start)"],
            ["QAOA calibrated estimate",
             "raw_estimate × estimate_multiplier"],
            ["Classical raw estimate",
             "safety_factor × (parse_overhead + result_overhead "
             "+ α_qubit × n + α_candidate × state_count)"],
            ["Calibrated coefficients",
             "safety_factor = 2.0,  α_qaoa = 5×10⁻⁷,  estimate_multiplier = 2.75,  "
             "α_classical = 3×10⁻⁶,  α_qubit = 0.01"],
        ],
        [4.2, 12.2],
    )
    _dark_callout(
        doc,
        "Live ETA.",
        "The pre-run estimate is based on algorithm complexity before execution. "
        "The live backend ETA during a job reflects actual measured speed and is "
        "significantly more accurate. Use the pre-run estimate for go/no-go decisions; "
        "watch the live ETA for real-time progress.",
        C["cyan"],
    )

    doc.add_page_break()

    # ── 7. Execution ───────────────────────────────────────────────────────
    _section_banner(doc, "7.  Execution and Live Progress")
    doc.add_heading("7.1  Job Architecture", level=2)
    doc.add_paragraph(
        "Runs are submitted as asynchronous backend jobs executed on Google Cloud Run. "
        "The browser polls the backend for progress, logs, ETA, and result availability. "
        "The page does not need to remain focused during execution."
    )
    doc.add_heading("7.2  Client Log", level=2)
    doc.add_paragraph(
        "The Client Log shows timestamped messages from the frontend, including file "
        "load confirmation, status updates, and any error messages returned by the API."
    )
    doc.add_heading("7.3  Backend Optimization Log", level=2)
    doc.add_paragraph(
        "The Backend Optimization Log streams live messages from the worker job. "
        "Messages are shown newest-first and cover:"
    )
    _bullets(doc, [
        "Input validation and workbook parsing",
        "QUBO construction (including type constraint injection and constant delta)",
        "Classical candidate search and export",
        "QAOA configuration, angle initialization, and per-iteration progress",
        "Optimizer convergence and best-energy reporting per restart",
        "IBM hardware submission and status updates (if selected)",
        "Result table generation and finalization",
        "Export diagnostic summary (cap reasons, state counts)",
    ])
    _callout(
        doc,
        "Example log excerpt (16-qubit demo run).",
        "QAOA sample count: 5000\n"
        "Generating Version 6.1 result tables.\n"
        "Classical export requested 5016 rows; exported 9 rows; reason: "
        "unique_candidate_count_after_duplicate_removal_or_search_convergence\n"
        "QAOA export requested 5000 rows; exported 5000 rows; reason: "
        "requested_rows_within_safety_cap\n"
        "IBM hardware submitted on ibm_fez: job d81jg0ugbeec73akljpg (4096 shots).\n"
        "IBM hardware completed on ibm_fez (job d81jg0ugbeec73akljpg).",
        "F1F5F9", C["slate"],
    )

    # ── 8. Outputs ─────────────────────────────────────────────────────────
    _section_banner(doc, "8.  Outputs and Interpretation")
    _img(doc, results)
    doc.add_heading("8.1  Result Summary Cards", level=2)
    doc.add_paragraph(
        "Three summary cards appear side-by-side when all configured paths complete. "
        "Each card shows the best-found portfolio from that solver."
    )
    _table(
        doc,
        ["Card", "Source", "Content"],
        [
            ["Classical Result Summary", "Classical_Candidates",
             "Best bitstring, QUBO value, selected amount, budget gap, return, "
             "volatility, Sharpe ratio."],
            ["Quantum Result Summary", "QAOA_Best_QUBO",
             "Best QAOA bitstring by QUBO value, plus readout probability."],
            ["Quantum Result Summary (2nd opinion)", "IBM Hardware or Qiskit",
             "Best bitstring from hardware or Qiskit simulation, plus "
             "hardware/simulation probability."],
        ],
        [4.0, 4.4, 8.0],
    )
    _callout(
        doc,
        "Example (16-qubit demo, 2 layers, 100 iterations, 3 restarts).",
        "All three solvers converge on the same bitstring 0000000010011111 with "
        "QUBO = 0.794121, selected amount = 2,999,905.30 (budget gap −94.70), "
        "return 0.598, volatility 0.222, Sharpe 2.517. QAOA probability 0.004787; "
        "IBM hardware probability 0.001953.",
        "E0F2FE", C["navy"],
    )
    doc.add_heading("8.2  Reading the Metrics", level=2)
    _table(
        doc,
        ["Metric", "Definition", "Interpretation"],
        [
            ["Return", "Portfolio-weighted expected return proxy from Assets sheet.",
             "Higher is better, subject to risk and budget constraints."],
            ["Volatility", "Portfolio variance from AnnualizedCovariance matrix and "
             "selected weights.",
             "Lower is better; reflects diversification and individual asset risk."],
            ["Sharpe-like ratio", "(Return − risk-free rate) / Volatility.",
             "Higher is better. An indicator for comparison, not a financial guarantee."],
            ["Budget gap", "Selected amount − total budget.",
             "Small negative is typical. Zero = exact budget. Positive = over budget."],
            ["Budget-normalized metrics", "Rescaled to full budget assuming uninvested "
             "cash earns the risk-free rate.",
             "Comparable across runs with different budget utilization."],
            ["QUBO value", "Objective value of the bitstring under the full QUBO.",
             "Lower is better within the same penalty configuration."],
            ["Probability", "Readout probability of the bitstring from QAOA/hardware.",
             "Higher = circuit more likely to measure that portfolio."],
        ],
        [3.0, 5.6, 7.8],
    )
    doc.add_heading("8.3  QUBO Breakdown", level=2)
    doc.add_paragraph(
        "The QUBO breakdown decomposes the objective value into its constituent "
        "terms for the classical and quantum best portfolios."
    )
    _table(
        doc,
        ["Term", "Example value (16-qubit demo)", "What it means"],
        [
            ["Return term", "−0.09573",
             "Negative — good portfolios have large return terms lowering QUBO."],
            ["Risk term", "0.63309",
             "Dominates at default lambda_variance = 6. Reduce risk lambda to "
             "allow higher-return selections."],
            ["Budget term", "0.00000",
             "Near-zero means total budget constraint is approximately satisfied."],
            ["Type-budget term", "0.25676",
             "Enterprise Rails at +7.2% deviation drives this. Increase subtype "
             "penalty to enforce tighter adherence."],
            ["QUBO total", "0.79412", "Sum of all four terms."],
        ],
        [3.4, 4.4, 8.6],
    )
    doc.add_heading("8.4  Type-Budget Achievements", level=2)
    _table(
        doc,
        ["Column", "Description"],
        [
            ["Achieved", "Sum of fixed + selected variable exposures for this type."],
            ["Normalized", "Achieved / Budget. 1.000 = exact target hit."],
            ["Deviation", "Achieved − Budget (absolute)."],
            ["Relative deviation", "(Achieved / Budget) − 1.0 (percentage)."],
            ["Penalty contribution", "λ_k × (normalized − 1)²."],
        ],
        [3.8, 12.6],
    )
    _dark_callout(
        doc,
        "Penalty tuning.",
        "A large relative deviation with a small penalty contribution signals an "
        "underpowered penalty. Increase λ_k to enforce tighter adherence. If one "
        "subtype penalty dominates the full QUBO breakdown, consider whether its "
        "target is achievable given the available variable blocks.",
        C["amber"],
    )
    doc.add_heading("8.5  Top Candidates Tables", level=2)
    _bullets(doc, [
        "Top Classical Candidates: ranked by QUBO ascending, deduplicated. "
        "Source: heuristic search.",
        "Top Quantum Candidates by QUBO: QAOA results ranked by QUBO value, "
        "includes readout probability per candidate.",
        "Top Quantum Candidates (2nd opinion): IBM hardware or Qiskit simulation "
        "results ranked by QUBO.",
        "Top Quantum Samples by Probability: IBM hardware results sorted by "
        "observed hit count — shows which states the hardware most frequently "
        "measured, independent of QUBO quality.",
    ])
    _callout(
        doc,
        "Key insight.",
        "In a well-converged run, top-QUBO and top-probability candidates overlap "
        "substantially. Significant divergence indicates insufficient circuit depth, "
        "penalty imbalance, or too few restarts/iterations — useful diagnostic "
        "information, not a failure mode.",
        "ECFEFF", C["navy"],
    )
    doc.add_heading("8.6  Solver Comparison", level=2)
    doc.add_paragraph(
        "Side-by-side bars comparing classical, QAOA, and second-opinion results on "
        "QUBO value, portfolio return, volatility, and Sharpe ratio. Convergence "
        "across all three is the ideal outcome. Informative divergence should prompt "
        "investigation of settings rather than dismissal of results."
    )
    doc.add_heading("8.7  Optimization History", level=2)
    doc.add_paragraph(
        "Shows the QAOA expected QUBO energy as a function of optimizer iteration, "
        "across all restarts. Look for:"
    )
    _bullets(doc, [
        "Smooth descent without early stagnation — stagnation may indicate too "
        "few iterations or a difficult landscape at this depth.",
        "Multiple restarts converging to the same energy — strong consistency is a "
        "positive sign of a well-defined optimum.",
        "Large energy gap between restarts — indicates sensitivity to initialization; "
        "consider more restarts or enabling warm start.",
    ])

    doc.add_page_break()

    # ── 9. IBM Hardware ────────────────────────────────────────────────────
    _section_banner(doc, "9.  IBM Quantum Hardware Second Opinion")
    doc.add_heading("9.1  Overview", level=2)
    doc.add_paragraph(
        "When Qiskit on IBM Hardware is selected, the backend reconstructs the "
        "optimized QAOA circuit in Qiskit, transpiles it for a real IBM quantum "
        "device, submits the job, and integrates hardware measurement counts into the "
        "result. This enables cross-platform validation and demonstrates end-to-end "
        "quantum hardware execution within the same workflow."
    )
    doc.add_heading("9.2  IBM Quantum Session Setup", level=2)
    _table(
        doc,
        ["Field", "Description"],
        [
            ["IBM API token",
             "Personal IBM Quantum API token. Used only for this run; "
             "not stored in result files or transmitted beyond the backend."],
            ["IBM instance",
             "IBM Quantum platform instance string (default: open-instance)."],
            ["Optional backend override",
             "Leave blank for auto-selection of the best available backend. "
             "Specify a backend name (e.g., ibm_fez) to override."],
            ["Fractional gates",
             "When enabled, uses RZZ gates instead of CX-RZ-CX for ZZ interactions. "
             "May reduce transpiled depth on hardware supporting native RZZ."],
            ["Parallelization",
             "Groups non-overlapping two-qubit terms into parallel rounds, "
             "potentially reducing circuit depth."],
        ],
        [4.2, 12.2],
    )
    doc.add_heading("9.3  Pre-Run Hardware Depth Estimate", level=2)
    _table(
        doc,
        ["Metric", "Description", "Example (16q, p=2)"],
        [
            ["Est. total gates", "Logical gate count before transpilation.", "800"],
            ["Est. 2Q gates", "Logical CX / ZZ gate count.", "480"],
            ["Est. sequential 2Q", "Depth counting only 2-qubit operations — "
             "the primary hardware feasibility metric.", "480"],
            ["HW reference limit", "Current reference threshold (ibm_fez).", "2,000"],
        ],
        [3.6, 7.6, 5.2],
    )
    _dark_callout(
        doc,
        "Transpilation scaling.",
        "Logical circuit gates expand substantially after transpilation to device-native "
        "gates. For 16 qubits / 2 layers, the 800 logical gates become 4,831 after "
        "transpilation to ibm_fez (CZ + RZ + SX basis), with a transpiled depth of 963 "
        "but only 283 sequential 2Q depth. Always check the hardware diagnostics panel "
        "after the run.",
        C["cyan"],
    )
    doc.add_heading("9.4  Shot Count Logic", level=2)
    _table(
        doc,
        ["Internal simulator mode", "Hardware shot count", "Reason"],
        [
            ["Exact probability (Lightning ≤ 24q)", "4,096 shots (default)",
             "Makes hardware results comparable in scale to the exact distribution."],
            ["Sampling mode", "Matches qaoa_shots setting",
             "Consistent sample count between simulator and hardware."],
        ],
        [4.8, 3.6, 8.0],
    )
    doc.add_heading("9.5  Bit Order and Counts Decoding", level=2)
    doc.add_paragraph(
        "IBM Qiskit uses a reversed bit ordering convention: the Qiskit counts key "
        "cN-1…c0 maps to optimizer order q0…qN-1. The backend handles this reversal "
        "automatically. All candidate tables report bitstrings in consistent optimizer "
        "order (q0 = first variable block, qN-1 = last)."
    )
    doc.add_heading("9.6  IBM Hardware Diagnostics (Full Response)", level=2)
    _table(
        doc,
        ["Diagnostic", "Description"],
        [
            ["Backend", "Selected IBM device name (e.g., ibm_fez)."],
            ["IBM job ID", "Job identifier for audit or IBM Quantum dashboard lookup."],
            ["Qiskit / Runtime version", "SDK versions used (e.g., Qiskit 2.4.1 / Runtime 0.46.1)."],
            ["Queue wait / Execution / Total", "Actual hardware timing breakdown."],
            ["Transpiled depth", "Circuit depth on target device after transpilation."],
            ["Transpiled 2Q gates", "Number of native 2-qubit gates after transpilation."],
            ["Transpiled 2Q depth", "Sequential 2Q depth after transpilation."],
            ["Parse status", "Confirms measurement counts were successfully decoded."],
        ],
        [4.2, 12.2],
    )

    doc.add_page_break()

    # ── 10. Review files & code exports ────────────────────────────────────
    _section_banner(doc, "10.  Review Files and Code Exports")
    _img(doc, exports)
    doc.add_heading("10.1  Review File (JSON)", level=2)
    _bullets(doc, [
        "Save Review File: captures the full result view (metrics, tables, charts, "
        "logs, settings, workbook filename) as a self-contained local JSON file.",
        "Download Raw JSON Data: saves the raw backend result JSON, including the "
        "embedded code-export package. Primary artifact for audit, archival, and "
        "downstream processing.",
        "Load Review / Raw JSON File: restores a previously saved result. A review "
        "file loaded without the original workbook cannot be rerun — it is a "
        "read-only result snapshot.",
    ])
    doc.add_heading("10.2  Code Exports", level=2)
    doc.add_paragraph(
        "Code exports generate executable quantum circuit files using the optimized "
        "QAOA angles (best_gammas and best_betas) and QUBO-derived Ising Hamiltonian "
        "embedded in the result JSON."
    )
    _table(
        doc,
        ["Export", "Format", "Description", "Level required"],
        [
            ["Qiskit Notebook", ".ipynb",
             "Builds QuantumCircuit, runs statevector simulation, computes "
             "probabilities. Includes optional IBM Runtime cell for hardware submission.",
             "Tester+"],
            ["Qiskit Python", ".py",
             "Script-style reproducible review for CI or code-review workflows.",
             "Tester+"],
            ["PennyLane Notebook", ".ipynb",
             "Cross-framework reconstruction using default.qubit. Python 3.11+ "
             "kernel recommended.",
             "Tester+"],
            ["Google Cirq Notebook", ".ipynb",
             "Alternative circuit framework for Cirq simulator comparison.",
             "Tester+"],
            ["Quantinuum Notebook", ".ipynb",
             "Planned for a later release.",
             "TBD"],
        ],
        [3.4, 1.6, 7.2, 3.2],
    )
    _dark_callout(
        doc,
        "Export availability.",
        "Code exports require a completed QAOA run or a loaded result JSON containing "
        "the code-export package. Buttons are disabled until this package is available. "
        "Code exports require a Tester-level key or above.",
        C["cyan"],
    )
    doc.add_heading("10.3  Circuit Construction in Exports", level=2)
    doc.add_paragraph(
        "All exported circuits use the same gate decomposition as the internal "
        "PennyLane circuit. The Qiskit circuit standard construction is:"
    )
    _bullets(doc, [
        "Initialization: H gate on every qubit.",
        "Cost layer (per p): RZ(2γ·h_i) for each single-qubit Ising term; "
        "CX → RZ(2γ·J_ij) → CX for each two-qubit ZZ interaction term.",
        "Mixer layer (per p): RX(2β) on every qubit.",
        "Measurement: measure_all() appended for sampler-based execution.",
    ])
    doc.add_paragraph(
        "The circuit metadata includes the optimizer bitstring order (q0…qN-1) and "
        "the Qiskit counts key order (cN-1…c0) so that any consumer can apply the "
        "correct reversal independently."
    )

    doc.add_page_break()

    # ── 11. Access levels ──────────────────────────────────────────────────
    _section_banner(doc, "11.  Access Levels and Limits")
    _img(doc, access)
    doc.add_paragraph(
        "Access is controlled by license key, validated at submission time. "
        "Exact limits are defined in the backend configuration and may evolve."
    )
    _table(
        doc,
        ["Level", "ID", "Max qubits", "Lightning Sim limits", "Max runtime", "Upload"],
        [
            ["Public Demo",    "0", "8",  "8q / 1 layer / 10 iter / 1 restart",    "60 s",   "5 MB"],
            ["Qualified Demo", "1", "16", "8q / 2 layers / 20 iter / 1 restart",   "4 min",  "10 MB"],
            ["Tester",         "2", "24", "16q / 6 layers / 200 iter / 3 restarts","45 min", "25 MB"],
            ["Internal Power", "3", "24", "24q / 8 layers / 300 iter / 5 restarts","2 h",    "25 MB"],
            ["Int. QAOA 30",   "4", "30", "30q / 6 layers / 300 iter / 3 restarts","2 h",    "25 MB"],
            ["Internal Ultra", "5", "35", "35q / 10 layers / 300 iter / 3 restarts","30 d",  "25 MB"],
        ],
        [3.2, 0.9, 2.2, 6.2, 2.2, 1.9],
    )
    doc.add_heading("Feature availability by level:", level=3)
    _table(
        doc,
        ["Feature", "Public Demo", "Qualified Demo", "Tester", "Internal"],
        [
            ["Classical mode",           "✓", "✓", "✓", "✓"],
            ["QAOA Lightning Sim",       "✓", "✓", "✓", "✓"],
            ["QAOA Tensor Sim",          "—", "—", "✓", "✓"],
            ["Qiskit simulation (2nd op)","—", "—", "✓", "✓"],
            ["IBM Hardware (2nd op)",    "—", "—", "✓", "✓"],
            ["Code exports",             "—", "—", "✓", "✓"],
            ["Medium/Large worker",      "—", "—", "✓*","✓"],
        ],
        [4.8, 2.8, 3.4, 2.6, 2.8],
    )
    doc.add_paragraph("* Medium available at Tester; Large at Internal Power+.")
    _callout(
        doc,
        "Client note.",
        "If a button or mode is disabled, it usually means the current key level "
        "does not allow that feature, the result is not yet completed, or the loaded "
        "review file does not contain the required code-export package.",
        "E0F2FE", C["navy"],
    )

    doc.add_page_break()

    # ── 12. Worker profiles ────────────────────────────────────────────────
    _section_banner(doc, "12.  Worker Profiles")
    doc.add_paragraph(
        "Worker profiles control the backend CPU and memory resources allocated to "
        "the Cloud Run job. Each run is dispatched to a profile-matched worker."
    )
    _table(
        doc,
        ["Profile", "vCPU", "Memory", "QAOA qubit capacity", "Level required", "Description"],
        [
            ["Small",  "2", "2 GiB", "Up to ~18 qubits", "All levels",    "For small examples and quick tests."],
            ["Medium", "4", "4 GiB", "Up to ~25 qubits", "Tester+",       "For larger QAOA simulations."],
            ["Large",  "4", "8 GiB", "Policy-gated",     "Internal Power+","For heavy runs at the upper qubit limits."],
        ],
        [2.2, 1.4, 2.2, 3.6, 3.2, 4.0],
    )
    _bullets(doc, [
        "Automatic capacity check: if the selected profile is too small for the problem "
        "size and mode, the backend returns an error before the run starts, specifying "
        "the minimum required profile.",
        "Profile selection: 16-qubit QAOA Lightning Sim fits comfortably on Small "
        "(peak memory ~0.49 GiB observed in the 16-qubit demo). Use Medium for "
        "20–24 qubit simulations or many restarts at high depth.",
        "Large is reserved for the heaviest experiments at or near the upper qubit "
        "limits per key level.",
    ])

    # ── 13. Memory telemetry ───────────────────────────────────────────────
    _section_banner(doc, "13.  Memory Telemetry")
    doc.add_paragraph(
        "The backend worker reports memory usage observations at regular intervals "
        "(approximately every 10 seconds during execution). Memory is read from Linux "
        "cgroup counters inside the container. The Memory Diagnostics panel (full "
        "response level) shows:"
    )
    _table(
        doc,
        ["Metric", "Description"],
        [
            ["Current memory usage", "GiB used at last observation."],
            ["Memory limit", "Profile memory ceiling (2 / 4 / 8 GiB)."],
            ["Remaining within job limit", "Headroom at last sample."],
            ["Peak memory usage", "Maximum observed across all samples during the run."],
            ["Sample count", "Number of memory observations recorded."],
            ["Memory usage chart", "Time-series plot against the job-limit ceiling."],
        ],
        [4.6, 11.8],
    )
    _dark_callout(
        doc,
        "Planning profiles.",
        "Peak memory is the most useful figure for profile selection. The 16-qubit "
        "demo peaked at 0.49 GiB on a 2 GiB Small worker (75% headroom). Expect "
        "memory to scale roughly as O(2^n) for Lightning Sim exact-probability mode.",
        C["green"],
    )

    doc.add_page_break()

    # ── 14. Good practice & troubleshooting ────────────────────────────────
    _section_banner(doc, "14.  Good Practice and Troubleshooting")
    doc.add_heading("14.1  Recommended Workflow Discipline", level=2)
    _bullets(doc, [
        "Always inspect the workbook before running. The inspect step is fast and "
        "catches structural errors that would waste a run.",
        "Use Classical Only mode first on a new workbook to validate inputs and see "
        "the classical baseline before committing a QAOA run.",
        "Save a Review File for any run that may be discussed with clients, technical "
        "reviewers, or stakeholders. It is the complete reproducibility artifact.",
        "Treat QAOA results as rapid prototypes, not production investment decisions. "
        "The tool is designed for education, demonstration, and structured exploration.",
        "Compare top-QUBO and top-probability tables deliberately. Disagreement "
        "between them is diagnostic information, not a bug.",
        "Use the QUBO breakdown panel to diagnose penalty balance before adjusting "
        "settings.",
    ])
    doc.add_heading("14.2  Penalty Tuning Guide", level=2)
    _table(
        doc,
        ["Parameter", "Starting value", "Increase when…", "Decrease when…"],
        [
            ["Budget lambda (λ_b)", "50",
             "Selected amount is consistently far from budget.",
             "Budget term dominates QUBO and all candidates look the same."],
            ["Risk lambda (λ_v)", "6",
             "You want lower-volatility portfolios.",
             "High-return selections are being suppressed excessively."],
            ["Subtype penalty (λ_k)", "50",
             "A subtype consistently misses its target.",
             "One subtype penalty dominates the full QUBO breakdown."],
        ],
        [3.4, 2.8, 5.0, 5.2],
    )
    doc.add_heading("14.3  Troubleshooting Reference", level=2)
    _table(
        doc,
        ["Situation", "Recommended action"],
        [
            ["Workbook does not inspect",
             "Check for required sheets (Settings, Assets, AnnualizedCovariance). "
             "Ensure required columns are present and headers match exactly. "
             "Verify numeric fields contain only numbers."],
            ["Workbook inspects but shows warnings",
             "Review warnings in the Workbook Summary. Common causes: fixed "
             "holdings exceed budget; no variable blocks defined."],
            ["Type constraint validation fails",
             "Ensure Additional Type Constraints is an integer 0–5. Check that "
             "Type X Size columns exist in Assets with only numeric values. "
             "Verify Type X Budget is positive and penalty is ≥ 0."],
            ["Runtime estimate exceeds limit",
             "Reduce qubits, QAOA layers, iterations, or restarts. Upgrade key "
             "level or worker profile."],
            ["QAOA result differs significantly from classical",
             "Expected for short circuits. Increase layers, iterations, restarts. "
             "Enable warm start. Review penalty balance."],
            ["All QAOA candidates have similar bitstrings",
             "Circuit may have converged prematurely. Increase restarts, enable "
             "restart perturbation, or increase iterations."],
            ["IBM hardware job fails or times out",
             "Verify IBM token is valid. Review pre-run circuit depth estimate — "
             "if sequential 2Q depth approaches 2,000, reduce layers."],
            ["PennyLane notebook import error",
             "Use a Python 3.11+ Jupyter kernel and run the first install cell "
             "(%pip install pennylane), then restart the kernel."],
            ["Qiskit notebook cell fails",
             "Run the first install cell (%pip install qiskit qiskit-aer "
             "qiskit-ibm-runtime), restart, and rerun from the top."],
            ["Code export buttons are disabled",
             "Complete a QAOA run or load a result JSON containing the "
             "code-export package. Ensure key level is Tester or above."],
            ["Memory exceeded during job",
             "Upgrade the worker profile or reduce problem size "
             "(fewer qubits, fewer iterations/restarts)."],
        ],
        [5.2, 11.2],
    )

    doc.add_page_break()

    # ── Appendix A ─────────────────────────────────────────────────────────
    _section_banner(doc, "Appendix A  —  Settings Field Reference",
                    bg=C["slate"], fg=C["white"])
    doc.add_heading("A.1  Excel Settings Sheet", level=2)
    _table(
        doc,
        ["Field name (Settings sheet)", "Type", "Description"],
        [
            ["Budget", "Numeric", "Total portfolio budget in USD (or chosen currency)."],
            ["Lambda Budget", "Numeric", "Penalty weight for total budget constraint."],
            ["Lambda Variance", "Numeric", "Penalty weight for portfolio variance."],
            ["Risk Free Rate", "Numeric (decimal)", "Annual risk-free rate, e.g., 0.04 for 4%."],
            ["QAOA P", "Integer", "QAOA depth (layers)."],
            ["QAOA Maxiter", "Integer", "Maximum optimizer iterations per restart."],
            ["QAOA Multistart Restarts", "Integer", "Number of independent optimizer restarts."],
            ["QAOA Shots", "Integer", "Shots for sampling mode."],
            ["Warm Start", "Boolean", "true/false — enable layerwise warm start."],
            ["Restart Perturbation", "Numeric (decimal)", "Noise magnitude for restart initialization."],
            ["Random Seed", "Integer 0–4,294,967,295", "Optional reproducibility seed."],
            ["Additional Type Constraints", "Integer 0–5", "Number of active subtype budget constraints."],
            ["Type A Name … Type E Name", "Text", "User-facing label for each subtype."],
            ["Type A Budget … Type E Budget", "Numeric", "Target exposure for each subtype."],
            ["Type A Budget Penalty … Type E Budget Penalty", "Numeric ≥ 0",
             "Penalty weight for each subtype constraint."],
        ],
        [5.6, 3.2, 7.6],
    )
    doc.add_heading("A.2  UI Form Fields", level=2)
    _table(
        doc,
        ["UI field", "Type", "Notes"],
        [
            ["License key", "Text", "Optional for public demo; required for keyed tiers and code exports."],
            ["Excel file", ".xlsx", "Use a demo workbook as a template for required structure."],
            ["Mode", "Dropdown", "classical_only / qaoa_lightning_sim / qaoa_tensor_sim."],
            ["2nd opinion", "Dropdown", "internal_only / qiskit_export / ibm_external_run."],
            ["Worker profile", "Dropdown", "small / medium / large. Availability depends on key level."],
            ["Response level", "Dropdown", "compact / standard / full."],
            ["Layers", "Integer", "QAOA depth p."],
            ["Iterations", "Integer", "Optimizer iteration budget."],
            ["Restarts", "Integer", "Independent optimizer starts."],
            ["QAOA shots", "Integer or exact", "exact shown in exact-probability mode."],
            ["Budget lambda", "Numeric", "Overrides workbook Lambda Budget if set."],
            ["Risk lambda", "Numeric", "Overrides workbook Lambda Variance if set."],
            ["Risk-free rate", "Numeric", "Overrides workbook Risk Free Rate if set."],
            ["Random seed", "Integer", "Optional. Overrides workbook value if set."],
            ["IBM API token", "Text", "Session-only. Not stored in result files."],
            ["IBM instance", "Text", "IBM Quantum platform instance (default: open-instance)."],
            ["IBM backend override", "Text", "Leave blank for auto-select."],
        ],
        [3.8, 2.4, 10.2],
    )

    doc.add_page_break()

    # ── Appendix B ─────────────────────────────────────────────────────────
    _section_banner(doc, "Appendix B  —  Glossary",
                    bg=C["slate"], fg=C["white"])
    _table(
        doc,
        ["Term", "Plain-language meaning"],
        [
            ["QUBO",
             "Quadratic Unconstrained Binary Optimization. A problem formulation "
             "where all decisions are binary (0/1) and the objective is a quadratic "
             "polynomial over those variables."],
            ["Ising model",
             "An equivalent quantum-friendly reformulation of the QUBO. Binary "
             "variables become ±1 spin variables; the QUBO matrix maps to "
             "Pauli Z (σ_z) operators."],
            ["QAOA",
             "Quantum Approximate Optimization Algorithm. A variational quantum "
             "algorithm with alternating cost and mixer unitaries, using classically "
             "optimized angle parameters."],
            ["PennyLane",
             "Quantum ML framework (Xanadu) used as the V9 simulation backend. "
             "lightning.qubit = C++ statevector; default.tensor = tensor network."],
            ["Statevector",
             "Complete quantum state as a vector of 2^n complex amplitudes. "
             "Squared magnitudes give exact bitstring probabilities."],
            ["Bitstring",
             "A string of n binary digits (0/1) representing a portfolio selection. "
             "Bit i = 1 means variable block i is selected."],
            ["Fixed asset",
             "Decision Role = fixed. Always included; no qubit consumed."],
            ["Variable asset",
             "Decision Role = variable. Binary decision; consumes one qubit."],
            ["QUBO value",
             "The numerical objective value assigned to a bitstring. Lower is better "
             "within the same penalty configuration."],
            ["Exact probability mode",
             "Lightning Sim ≤ 24 qubits. Full probability distribution over all 2^n "
             "bitstrings computed without sampling noise."],
            ["Sampling mode",
             "Measurement-based mode. Circuit measured qaoa_shots times; "
             "probabilities estimated from observed frequencies."],
            ["Second opinion",
             "Independent reconstruction in Qiskit or on IBM hardware for "
             "cross-framework validation."],
            ["Review file",
             "JSON snapshot of a completed result. Restores full view without "
             "rerunning the backend."],
            ["Code-export package",
             "Self-contained data block (QUBO, Ising terms, optimized angles) "
             "embedded in result JSON for regenerating notebooks."],
            ["Worker profile",
             "Cloud Run job resource tier (Small / Medium / Large)."],
            ["Runtime estimate",
             "Pre-run predicted job duration. Live backend ETA during execution "
             "is more accurate."],
            ["Budget gap",
             "Selected amount − total budget. Small negative is typical."],
            ["Type-budget term",
             "Sum of all active subtype budget penalty contributions in the QUBO."],
            ["Warm start",
             "Layer p angles seeded from optimized layer p−1 solution."],
            ["Transpilation",
             "Converting a logical circuit to device-native gates, typically "
             "increasing gate count and depth."],
            ["Sequential 2Q depth",
             "Critical-path length counting only two-qubit gates — the primary "
             "hardware execution feasibility metric."],
        ],
        [3.6, 12.8],
    )

    # ── Disclaimer ─────────────────────────────────────────────────────────
    _section_banner(doc, "Disclaimer", bg=C["slate"], fg=C["muted"])
    doc.add_paragraph(
        "This manual describes a rapid prototyping and education tool. Outputs are "
        "intended for technical review, experimentation, demonstration, and structured "
        "discussion. They are not financial advice and should not be used as the sole "
        "basis for investment, risk, or production technology decisions. Quantum "
        "hardware results include noise effects intrinsic to current-generation "
        "devices and should be interpreted accordingly."
    )

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_doc()
